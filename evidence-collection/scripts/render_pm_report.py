#!/usr/bin/env python3
"""Render a project-manager-first heavy-case evidence report.

The input is a case-local snapshot containing:
  results, attachments, pm_status, run_state, action_queue.

The renderer deliberately separates:
  - AI checked and found a lead;
  - AI checked but found no result in the stated scope;
  - not checked because login is required;
  - not checked because a human/offline step is required;
  - attempted but failed and still needs action.

The DOCX is generated and validated in one pass. It is never post-processed
after validation.
"""

from __future__ import annotations

import urllib.parse
import argparse
import hashlib
import json
import os
import re
import tempfile
import zipfile
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import quote

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import (
    WD_ALIGN_VERTICAL,
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SCHEMA_VERSION = "pm-report-summary/1.0"
STATUS_ORDER = (
    "AI_CHECKED_FOUND",
    "AI_CHECKED_NO_RESULT",
    "NOT_CHECKED_LOGIN",
    "NOT_CHECKED_MANUAL",
    "RUN_FAILED_NEEDS_ACTION",
)
STATUS = {
    "AI_CHECKED_FOUND": {
        "label": "已查看｜有线索",
        "fill": "E6F4EE",
        "accent": "0F6B4D",
        "short": "有线索",
    },
    "AI_CHECKED_NO_RESULT": {
        "label": "已查看｜未发现",
        "fill": "E8F0F6",
        "accent": "365A7C",
        "short": "未发现",
    },
    "NOT_CHECKED_LOGIN": {
        "label": "待处理｜需登录",
        "fill": "FFF4CE",
        "accent": "8A6100",
        "short": "需登录",
    },
    "NOT_CHECKED_MANUAL": {
        "label": "待处理｜需人工",
        "fill": "FCEBD8",
        "accent": "9A4A00",
        "short": "需人工",
    },
    "RUN_FAILED_NEEDS_ACTION": {
        "label": "待处理｜运行失败",
        "fill": "FBE7EA",
        "accent": "A51D2D",
        "short": "运行失败",
    },
}

FONT = "Microsoft YaHei"
INK = "172B3A"
MUTED = "5C6B76"
NAVY = "12344A"
BLUE = "1D617A"
BORDER = "B9C7D0"
HEADER_FILL = "DDE8EE"
LIGHT_FILL = "F5F8FA"

# compact_reference_guide with an explicit A4-landscape PM matrix override.
PAGE_MARGIN_IN = 0.55
BODY_SIZE = 8.2
TABLE_SIZE = 7.3
CELL_MARGINS = {"top": 70, "bottom": 70, "start": 90, "end": 90}
MAIN_WIDTHS_CM = (2.50, 3.60, 4.10, 7.40, 5.30, 3.90)
PENDING_WIDTHS_CM = (1.25, 1.55, 2.65, 5.30, 8.00, 7.05)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Render a project-manager execution report from a run-first case snapshot."
    )
    parser.add_argument("--input", required=True, help="Combined case snapshot JSON")
    parser.add_argument("--output", required=True, help="Final DOCX")
    parser.add_argument("--summary-output", required=True, help="Normalized PM summary JSON")
    parser.add_argument("--title", default="重案证据清单｜项目经理执行版")
    parser.add_argument("--rights-holder", default="")
    parser.add_argument("--defendant-product", default="")
    parser.add_argument("--source-url", default="")
    return parser.parse_args()


def load_json(path: str | Path) -> Any:
    return json.loads(Path(path).read_text())


def list_field(obj: Any, keys: Iterable[str]) -> list[dict[str, Any]]:
    if isinstance(obj, list):
        return obj
    if isinstance(obj, dict):
        for key in keys:
            value = obj.get(key)
            if isinstance(value, list):
                return value
    return []


def compact(text: Any, limit: int = 240) -> str:
    if isinstance(text, list):
        text = "；".join(str(value) for value in text if value)
    text = re.sub(r"\s+", " ", str(text or "")).strip()
    if len(text) <= limit:
        return text
    return text[: max(1, limit - 1)].rstrip("，；。 ") + "…"


def names(items: list[dict[str, Any]], key: str = "platform", limit: int = 6) -> str:
    values: list[str] = []
    for item in items:
        value = str(item.get(key) or "").strip()
        if value and value not in values:
            values.append(value)
    if not values:
        return "无"
    if len(values) <= limit:
        return "、".join(values)
    return "、".join(values[:limit]) + f"等{len(values)}个平台"


def as_uri(path: str) -> str:
    if path.startswith(("http://", "https://", "file://")):
        return path
    try:
        return Path(path).resolve().as_uri()
    except Exception:
        return "file://" + quote(path)


def save_items_for(row: dict[str, Any], result: dict[str, Any]) -> str:
    status = row["manager_status"]
    corpus = " ".join(
        [
            str(row.get("objective") or ""),
            str(row.get("evidence_objectives") or ""),
        ]
    )
    if any(token in corpus for token in ("购买", "订单", "物流", "收款", "发票", "实物", "封样")):
        return "订单页、付款记录、发票、物流面单、商品实物、外包装各面照片和封样记录。"
    if any(token in corpus for token in ("合同", "授权", "经销", "代理", "加盟", "接触机会")):
        return "合同原件、签章页、附件、往来邮件/聊天、付款或发票，以及签约主体和时间线。"
    if any(token in corpus for token in ("问卷", "专家", "访谈", "消费者")):
        return "问卷原件、原始数据、访谈录音/纪要、签名确认和专家意见原件。"
    if any(token in corpus for token in ("展会", "线下", "仓库", "工厂", "走访")):
        return "现场全景和细节照片、地址定位、时间、人员说明、展商名录或主办方材料。"
    if any(token in corpus for token in ("残疾", "年龄", "精神状态", "个人状况")):
        return "仅保存合法取得且与案件必要相关的证明、裁判认定或当事人提交材料，并记录来源；不得非法收集隐私。"
    if any(token in corpus for token in ("财产", "失信", "执行人", "偿付", "破产")):
        return "查询结果页/无结果页截图、主体名称、检索词、查询时间、详情页PDF和来源网址。"
    if status == "NOT_CHECKED_LOGIN":
        return "登录后的结果页/无结果页截图、检索词、查询时间、详情页PDF和来源网址。"
    if status == "NOT_CHECKED_MANUAL":
        return "人工取得的原件或完整截图、取得方式、时间、提供人和来源说明。"
    if status == "RUN_FAILED_NEEDS_ACTION":
        return "重试后的结果页；若仍失败，保存错误提示、时间、平台、检索词和操作步骤。"
    if status == "AI_CHECKED_NO_RESULT":
        return "本次无结果页面或调用记录、检索平台、关键词和查询时间，供后续复核。"
    sources = result.get("sources") or []
    if sources:
        return "现有页面截图、来源链接、查询时间；作为诉讼证据时再补平台证明或公证保全。"
    return "来源链接、查询时间、完整页面截图或导出文件。"


def gate_run_state(run_state: dict[str, Any]) -> None:
    counts = (run_state.get("counts") or {}).get("by_execution_state") or {}
    blocked = {
        key: int(counts.get(key, 0) or 0)
        for key in ("READY", "RUNNING", "RETRYABLE")
        if int(counts.get(key, 0) or 0) > 0
    }
    if run_state.get("active_action_id"):
        blocked["active_action_id"] = run_state["active_action_id"]
    if blocked:
        raise ValueError(f"正式报告门禁未通过，仍有可执行动作：{blocked}")


def build_summary(snapshot: dict[str, Any]) -> dict[str, Any]:
    results = list_field(snapshot.get("results"), ())
    pm_raw = snapshot.get("pm_status") or {}
    pm_rows = list_field(pm_raw, ("rows", "items", "classifications", "results"))
    attachments = list_field(
        snapshot.get("attachments"),
        ("rows", "attachments", "items", "evidence_attachments"),
    )
    run_state = snapshot.get("run_state") or {}
    action_queue = snapshot.get("action_queue") or {}
    gate_run_state(run_state)

    if len(results) != 61 or len(pm_rows) != 61 or len(attachments) != 61:
        raise ValueError(
            f"61行门禁失败：results={len(results)}, pm={len(pm_rows)}, attachments={len(attachments)}"
        )
    result_by_id = {str(row.get("item_id") or row.get("row_id")): row for row in results}
    attachment_by_id = {str(row.get("row_id")): row for row in attachments}
    pm_ids = [str(row.get("row_id")) for row in pm_rows]
    if len(set(pm_ids)) != 61 or set(pm_ids) != set(result_by_id):
        raise ValueError("PM状态未唯一覆盖全部61行")
    if set(pm_ids) != set(attachment_by_id):
        raise ValueError("附件清单未唯一覆盖全部61行")

    status_counts = Counter(str(row.get("manager_status")) for row in pm_rows)
    unknown = set(status_counts) - set(STATUS_ORDER)
    if unknown:
        raise ValueError(f"未知项目经理状态：{sorted(unknown)}")

    normalized_rows: list[dict[str, Any]] = []
    missing_images: list[str] = []
    for row in pm_rows:
        row_id = str(row["row_id"])
        result = result_by_id[row_id]
        attachment = attachment_by_id[row_id]
        status = str(row["manager_status"])
        action_counts = row.get("action_counts") or {}
        ai_checked_actions = row.get("ai_checked_actions") or []
        not_checked_actions = row.get("not_checked_actions") or []
        failed_actions = row.get("failed_actions") or []
        checked_platforms = [
            {
                "platform": item.get("platform", ""),
                "url": verified_canonical_url(item),
            }
            for item in ai_checked_actions
        ]
        for action_item, platform_item in zip(ai_checked_actions, checked_platforms):
            platform_name = str(platform_item.get("platform") or "")
            if (
                any(label in platform_name for label in ("小红书", "微博"))
                and str(action_item.get("outcome") or "") == "FOUND"
                and not platform_item["url"]
            ):
                raise ValueError(
                    f"{platform_name} FOUND缺少Chrome/Playwright已验证规范页，拒绝生成报告"
                )
        pending_platforms = [
            {"platform": item.get("platform", ""), "url": item.get("url", "")}
            for item in (not_checked_actions + failed_actions)
        ]
        images = attachment.get("images") or []
        for image in images:
            path = str(image.get("path") or "")
            if path and not Path(path).is_file():
                missing_images.append(path)
        objectives = row.get("evidence_objectives") or []
        objective = compact(objectives[0] if objectives else row_id, 110)
        id_category = compact(objectives[0] if objectives else "未明示类别", 38)
        row_id_label = f"{row_id}｜{id_category}"
        finding = compact(
            (result.get("findings") or [result.get("compact_note") or row.get("plain_summary")])[0],
            320,
        )
        if status == "AI_CHECKED_NO_RESULT":
            finding = (
                compact(finding, 230)
                + " 在本次平台和关键词范围内未发现，不代表绝对不存在。"
            )
        elif status == "RUN_FAILED_NEEDS_ACTION":
            finding = (
                "AI已尝试运行，但没有完成该项；这不是“未发现”。失败情况："
                + compact(row.get("plain_summary"), 220)
            )
        elif status in ("NOT_CHECKED_LOGIN", "NOT_CHECKED_MANUAL"):
            checked_count = int(action_counts.get("AI_CHECKED_FOUND", 0) or 0) + int(
                action_counts.get("AI_CHECKED_NO_RESULT", 0) or 0
            )
            finding = (
                f"该项尚未查完。AI已完成{checked_count}个动作；"
                + compact(row.get("plain_summary"), 220)
            )
        if row.get("is_mixed_checked_and_unresolved"):
            finding += " 本行为混合状态：已有AI结果，但仍有关键动作未完成。"

        manager_action = compact(row.get("manager_action"), 360)
        if status.startswith("AI_CHECKED_") and not manager_action:
            manager_action = (
                "先保存现有结果；只有当该项属于案件关键事实时，再按名称变体或登录平台复核。"
            )

        sources: list[dict[str, str]] = []
        for source in result.get("sources") or []:
                    url = verified_canonical_url(source)
            if url:
                sources.append(
                    {
                        "label": str(source.get("title") or "来源"),
                        "url": url,
                    }
                )
        for source in attachment.get("source_urls") or []:
            url = verified_canonical_url(source)
            if url:
                sources.append(
                    {
                        "label": str(source.get("label") or "页面"),
                        "url": url,
                    }
                )

        normalized_rows.append(
            {
                "row_id": row_id,
                "row_id_label": row_id_label,
                "priority": str(row.get("priority") or "P2"),
                "manager_status": status,
                "status_label": STATUS[status]["label"],
                "objective": objective,
                "is_mixed_checked_and_unresolved": bool(
                    row.get("is_mixed_checked_and_unresolved")
                ),
                "action_counts": action_counts,
                "ai_checked_platforms": checked_platforms,
                "pending_platforms": pending_platforms,
                "plain_summary": compact(row.get("plain_summary"), 300),
                "finding": finding,
                "manager_action": manager_action or "按本行平台和证据事项安排人工复核。",
                "save_items": save_items_for(row, result),
                "sources": sources,
                "images": images,
            }
        )

    if missing_images:
        raise ValueError(f"存在缺失图片：{missing_images[:5]}")

    line_actions_total = sum(
        int((row.get("action_counts") or {}).get("total", 0) or 0) for row in pm_rows
    )
    ai_completed_actions = sum(
        int((row.get("action_counts") or {}).get("AI_CHECKED_FOUND", 0) or 0)
        + int((row.get("action_counts") or {}).get("AI_CHECKED_NO_RESULT", 0) or 0)
        for row in pm_rows
    )
    pending_actions = sum(
        int((row.get("action_counts") or {}).get("NOT_CHECKED_LOGIN", 0) or 0)
        + int((row.get("action_counts") or {}).get("NOT_CHECKED_MANUAL", 0) or 0)
        + int(
            (row.get("action_counts") or {}).get("RUN_FAILED_NEEDS_ACTION", 0)
            or 0
        )
        for row in pm_rows
    )
    action_total = len(action_queue.get("actions") or [])
    seed_actions = action_total - line_actions_total
    ai_checked_rows = status_counts["AI_CHECKED_FOUND"] + status_counts[
        "AI_CHECKED_NO_RESULT"
    ]
    manager_pending_rows = 61 - ai_checked_rows

    summary = {
        "schema_version": SCHEMA_VERSION,
        "case_id": str(run_state.get("case_id") or pm_raw.get("case_id") or ""),
        "run_id": str(run_state.get("run_id") or pm_raw.get("run_id") or ""),
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "row_count": 61,
        "dashboard": {
            "ai_checked_rows": ai_checked_rows,
            "manager_pending_rows": manager_pending_rows,
            "status_counts": {key: status_counts[key] for key in STATUS_ORDER},
            "line_actions_total": line_actions_total,
            "ai_completed_actions": ai_completed_actions,
            "pending_actions": pending_actions,
            "seed_actions": seed_actions,
        },
        "rows": normalized_rows,
    }
    validate_summary(summary)
    return summary


def validate_summary(summary: dict[str, Any]) -> None:
    if summary.get("schema_version") != SCHEMA_VERSION:
        raise ValueError("PM summary schema_version错误")
    rows = summary.get("rows") or []
    if len(rows) != 61 or len({row["row_id"] for row in rows}) != 61:
        raise ValueError("PM summary必须唯一覆盖61行")
    counts = Counter(row["manager_status"] for row in rows)
    dashboard = summary.get("dashboard") or {}
    expected = dashboard.get("status_counts") or {}
    if {key: counts[key] for key in STATUS_ORDER} != {
        key: int(expected.get(key, -1)) for key in STATUS_ORDER
    }:
        raise ValueError("PM summary五态计数不一致")
    if dashboard.get("ai_checked_rows") + dashboard.get("manager_pending_rows") != 61:
        raise ValueError("PM summary首页行数不等于61")
    if (
        dashboard.get("ai_completed_actions", 0)
        + dashboard.get("pending_actions", 0)
        != dashboard.get("line_actions_total", -1)
    ):
        raise ValueError("PM summary动作计数不守恒")
    for row in rows:
        if row["status_label"] != STATUS[row["manager_status"]]["label"]:
            raise ValueError(f"{row['row_id']}状态文字与状态码不匹配")
        if not row.get("manager_action") or not row.get("save_items"):
            raise ValueError(f"{row['row_id']}缺项目经理动作或保存物")
        if row["manager_status"] == "AI_CHECKED_NO_RESULT" and "不代表绝对不存在" not in row[
            "finding"
        ]:
            raise ValueError(f"{row['row_id']}未发现结论缺边界")
        if row["manager_status"] == "RUN_FAILED_NEEDS_ACTION" and "不是“未发现”" not in row[
            "finding"
        ]:
            raise ValueError(f"{row['row_id']}失败状态与未发现未明确区分")


def set_font(
    run,
    size: float = BODY_SIZE,
    color: str = INK,
    bold: bool | None = None,
    italic: bool = False,
) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    run.italic = italic


def set_paragraph(
    paragraph,
    before: float = 0,
    after: float = 2,
    line: float = 1.05,
    align=None,
    keep: bool = False,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align
    if keep:
        fmt.keep_with_next = True


def add_text(
    target,
    text: str,
    size: float = BODY_SIZE,
    color: str = INK,
    bold: bool = False,
    italic: bool = False,
    after: float = 2,
    align=None,
):
    paragraph = target.add_paragraph()
    set_paragraph(paragraph, after=after, align=align)
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def clear_cell(cell) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    set_paragraph(paragraph, after=0)


def cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, **kwargs: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin in ("top", "start", "bottom", "end"):
        if margin not in kwargs:
            continue
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(kwargs[margin]))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_fixed(table, widths_cm: Iterable[float], indent_dxa: int = 120) -> None:
    widths_cm = tuple(widths_cm)
    widths_dxa = [round(width / 2.54 * 1440) for width in widths_cm]
    total_dxa = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:w"), str(total_dxa))
    tbl_width.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid_cols = table._tbl.tblGrid.findall(qn("w:gridCol"))
    for index, width_dxa in enumerate(widths_dxa):
        if index < len(grid_cols):
            grid_cols[index].set(qn("w:w"), str(width_dxa))
    for row in table.rows:
        for index, width in enumerate(widths_cm):
            if index < len(row.cells):
                row.cells[index].width = Cm(width)
                tc_pr = row.cells[index]._tc.get_or_add_tcPr()
                tc_width = tc_pr.find(qn("w:tcW"))
                if tc_width is None:
                    tc_width = OxmlElement("w:tcW")
                    tc_pr.append(tc_width)
                tc_width.set(qn("w:w"), str(widths_dxa[index]))
                tc_width.set(qn("w:type"), "dxa")
                set_cell_margins(row.cells[index], **CELL_MARGINS)
                row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, text: str, url: str, color: str = "0563C1") -> None:
    if not url:
        return
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:eastAsia"), FONT)
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(r_fonts)
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    r_pr.append(run_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "14")
    r_pr.append(size)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_cell_paragraph(
    cell,
    text: str,
    size: float = TABLE_SIZE,
    color: str = INK,
    bold: bool = False,
    after: float = 1.5,
    align=None,
):
    paragraph = cell.add_paragraph()
    set_paragraph(paragraph, after=after, line=1.02, align=align)
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color, bold=bold)
    return paragraph


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.clear()
    set_paragraph(paragraph, after=0)
    run = paragraph.add_run("重案证据调查｜项目经理执行版")
    set_font(run, size=7.5, color=MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(fp, after=0)
    run = fp.add_run("项目经理先看状态和待办，再看证据明细")
    set_font(run, size=7, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Inches(PAGE_MARGIN_IN)
    section.bottom_margin = Inches(PAGE_MARGIN_IN)
    section.left_margin = Inches(PAGE_MARGIN_IN)
    section.right_margin = Inches(PAGE_MARGIN_IN)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(BODY_SIZE)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08
    add_header_footer(doc)


def verified_canonical_url(record: dict[str, Any]) -> str:
    trace = record.get("execution_trace") if isinstance(record, dict) else None
    verification = trace.get("verification") if isinstance(trace, dict) else None
    verification = verification if isinstance(verification, dict) else {}
    tool = str(record.get("verification_tool") or verification.get("tool") or "")
    verified_at = str(
        record.get("page_verified_at") or verification.get("verified_at") or ""
    )
    page_verified = (
        record.get("page_verified") is True
        or verification.get("page_verified") is True
    )
    public_page = (
        record.get("public_page") is True
        or verification.get("public_page") is True
    )
    if tool not in {"chrome-devtools", "playwright"} or not page_verified or not verified_at:
        return ""
    if tool == "playwright" and not public_page:
        return ""
    raw = str(
        record.get("canonical_url")
        or verification.get("canonical_url")
        or verification.get("final_url")
        or ""
    ).strip()
    try:
        parts = urllib.parse.urlsplit(raw)
    except ValueError:
        return ""
    if parts.scheme.lower() not in {"http", "https"} or not parts.hostname:
        return ""
    host = parts.hostname.lower()
    path = parts.path or "/"
    platform = str(
        record.get("platform")
        or record.get("name")
        or record.get("label")
        or ""
    )
    if "小红书" in platform:
        if not host.endswith("xiaohongshu.com") or "/search_result" in path:
            return ""
        if not any(marker in path for marker in ("/explore/", "/discovery/item/", "/user/profile/")):
            return ""
    if "微博" in platform:
        if host == "s.weibo.com" or "/search" in path.lower() or path == "/":
            return ""
        if not (
            host == "weibo.com"
            or host.endswith(".weibo.com")
            or host == "m.weibo.cn"
            or host.endswith(".weibo.cn")
        ):
            return ""
    query = [
        (key, value)
        for key, value in urllib.parse.parse_qsl(parts.query, keep_blank_values=True)
        if not key.lower().startswith("utm_")
        and key.lower() not in {"spm", "source", "share", "share_source", "track_id"}
    ]
    return urllib.parse.urlunsplit(
        (
            parts.scheme.lower(),
            parts.netloc,
            path,
            urllib.parse.urlencode(query),
            "",
        )
    )


def add_title_block(
    doc: Document,
    title: str,
    rights_holder: str,
    defendant_product: str,
    source_url: str,
) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, before=4, after=2)
    run = paragraph.add_run("AI预跑结果与项目经理待办")
    set_font(run, size=9, color=BLUE, bold=True)
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, after=3)
    run = paragraph.add_run(title)
    set_font(run, size=23, color=NAVY, bold=True)
    metadata = []
    if rights_holder:
        metadata.append(f"权利人：{rights_holder}")
    if defendant_product:
        metadata.append(f"被告商品：{defendant_product}")
    metadata.append("阅读顺序：先看首页数字 → 再做待办 → 最后查证据明细")
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, after=8)
    run = paragraph.add_run("  |  ".join(metadata))
    set_font(run, size=8.5, color=MUTED)
    if source_url:
        add_hyperlink(paragraph, "打开种子商品", source_url)


def add_dashboard(doc: Document, summary: dict[str, Any]) -> None:
    dashboard = summary["dashboard"]
    counts = dashboard["status_counts"]
    table = doc.add_table(rows=2, cols=5)
    set_table_fixed(table, (5.36, 5.36, 5.36, 5.36, 5.36), indent_dxa=120)
    top = table.rows[0].cells
    metrics = [
        (
            "AI已查看",
            dashboard["ai_checked_rows"],
            "E8F0F6",
            "1D617A",
            "这24行已有可读结论",
        ),
        (
            "项目经理待处理",
            dashboard["manager_pending_rows"],
            "FFF0E5",
            "A44F00",
            "这37行仍需人推进",
        ),
    ]
    top[0].merge(top[1])
    top[2].merge(top[4])
    for cell, metric in zip((top[0], top[2]), metrics):
        clear_cell(cell)
        cell_shading(cell, metric[2])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph(p, after=1)
        run = p.add_run(str(metric[1]))
        set_font(run, size=24, color=metric[3], bold=True)
        p = add_cell_paragraph(
            cell,
            metric[0],
            size=10,
            color=metric[3],
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        add_cell_paragraph(
            cell,
            metric[4],
            size=7.2,
            color=MUTED,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    bottom = table.rows[1].cells
    cards = [
        ("有线索", counts["AI_CHECKED_FOUND"], "E6F4EE", "0F6B4D"),
        ("限定范围未发现", counts["AI_CHECKED_NO_RESULT"], "E8F0F6", "365A7C"),
        ("需登录", counts["NOT_CHECKED_LOGIN"], "FFF4CE", "8A6100"),
        ("需人工", counts["NOT_CHECKED_MANUAL"], "FCEBD8", "9A4A00"),
        ("运行失败", counts["RUN_FAILED_NEEDS_ACTION"], "FBE7EA", "A51D2D"),
    ]
    for cell, card in zip(bottom, cards):
        clear_cell(cell)
        cell_shading(cell, card[2])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(card[1]))
        set_font(run, size=14, color=card[3], bold=True)
        add_cell_paragraph(
            cell,
            card[0],
            size=7.4,
            color=card[3],
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    add_text(
        doc,
        (
            f"动作口径：{dashboard['line_actions_total']}个行级调查动作中，"
            f"AI完成{dashboard['ai_completed_actions']}个，待推进{dashboard['pending_actions']}个；"
            f"另有{dashboard['seed_actions']}个种子固定动作。"
        ),
        size=8,
        color=MUTED,
        after=3,
    )
    add_text(
        doc,
        "怎么判断：蓝绿状态表示AI确实看过；黄橙红状态表示项目经理仍要处理。"
        "“未发现”只说明本次平台和关键词范围内没有结果；“运行失败”则表示该项没有完成，二者不能互换。",
        size=8.4,
        color=INK,
        bold=True,
        after=4,
    )
    add_text(
        doc,
        "混合行处理：有些证据事项AI查过一部分，但登录、线下或失败动作仍未完成。"
        "这类行全部进入“待处理”，明细中会同时列出已查部分和剩余部分。",
        size=8,
        color=MUTED,
        after=2,
    )


def pending_reason(row: dict[str, Any]) -> str:
    status = row["manager_status"]
    if status == "NOT_CHECKED_LOGIN":
        return "平台需要登录，AI尚未进入数据库完成查询。"
    if status == "NOT_CHECKED_MANUAL":
        return "该项依赖购买、公证、实物、线下或人工材料，AI不能代替。"
    if status == "RUN_FAILED_NEEDS_ACTION":
        return "AI已尝试，但浏览器、平台限制或运行错误导致未完成。"
    return ""


def add_pending_table(doc: Document, summary: dict[str, Any]) -> None:
    heading = add_text(
        doc,
        "项目经理待办｜先处理这37项",
        size=16,
        color=NAVY,
        bold=True,
        after=2,
    )
    heading.paragraph_format.page_break_before = True
    add_text(
        doc,
        "排序原则：P0先锁定经营主体、交易链和实物生产者；P1再补规模、恶意和知名度；P2补资产、历史和外围线索。",
        size=8,
        color=MUTED,
        after=5,
    )
    pending = [
        row
        for row in summary["rows"]
        if row["manager_status"]
        in ("NOT_CHECKED_LOGIN", "NOT_CHECKED_MANUAL", "RUN_FAILED_NEEDS_ACTION")
    ]
    priority_order = {"P0": 0, "P1": 1, "P2": 2}
    pending.sort(key=lambda row: (priority_order.get(row["priority"], 9), row["row_id"]))
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    set_table_fixed(table, PENDING_WIDTHS_CM)
    headers = ("优先", "编号", "状态", "为什么没完成", "去哪里 / 做什么", "要保存什么")
    for cell, header in zip(table.rows[0].cells, headers):
        clear_cell(cell)
        cell_shading(cell, HEADER_FILL)
        add_cell_paragraph(
            cell,
            header,
            size=7.4,
            color=NAVY,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    set_repeat_table_header(table.rows[0])
    for item in pending:
        row = table.add_row()
        set_table_fixed(table, PENDING_WIDTHS_CM)
        for cell in row.cells:
            clear_cell(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        style = STATUS[item["manager_status"]]
        cell_shading(row.cells[2], style["fill"])
        values = (
            item["priority"],
            item["row_id"],
            item["status_label"],
            pending_reason(item),
            item["manager_action"],
            item["save_items"],
        )
        for index, value in enumerate(values):
            add_cell_paragraph(
                row.cells[index],
                compact(value, 220 if index >= 3 else 90),
                size=7.0,
                color=style["accent"] if index == 2 else INK,
                bold=index in (0, 1, 2),
                align=WD_ALIGN_PARAGRAPH.CENTER if index < 3 else WD_ALIGN_PARAGRAPH.LEFT,
            )
        for platform in item.get("pending_platforms") or []:
            url = str(platform.get("url") or "")
            if url:
                p = add_cell_paragraph(row.cells[4], "", size=6.8, after=0)
                add_hyperlink(p, f"打开{platform.get('platform') or '平台'}", url)


def add_scope_cell(cell, item: dict[str, Any]) -> None:
    checked = names(item.get("ai_checked_platforms") or [])
    pending = names(item.get("pending_platforms") or [])
    add_cell_paragraph(cell, "AI已查", size=6.8, color=BLUE, bold=True, after=0)
    add_cell_paragraph(cell, checked, size=7.0, color=INK, after=2)
    if item["manager_status"] in (
        "NOT_CHECKED_LOGIN",
        "NOT_CHECKED_MANUAL",
        "RUN_FAILED_NEEDS_ACTION",
    ):
        add_cell_paragraph(cell, "仍未完成", size=6.8, color=STATUS[item["manager_status"]]["accent"], bold=True, after=0)
        add_cell_paragraph(cell, pending, size=7.0, color=INK, after=1)
    for platform in (item.get("ai_checked_platforms") or [])[:2]:
        url = str(platform.get("url") or "")
        if url:
            p = add_cell_paragraph(cell, "", size=6.7, after=0)
            add_hyperlink(p, f"打开{platform.get('platform') or '来源'}", url)


def image_hash(path: str) -> str:
    digest = hashlib.sha256()
    with open(path, "rb") as handle:
        for chunk in iter(lambda: handle.read(65536), b""):
            digest.update(chunk)
    return digest.hexdigest()


def add_images(
    cell,
    item: dict[str, Any],
    seen: dict[str, str],
) -> int:
    images = item.get("images") or []
    selected: list[dict[str, Any]] = []
    repeated: list[tuple[dict[str, Any], str]] = []
    for image in images:
        path = str(image.get("path") or "")
        if not path or not Path(path).is_file():
            continue
        digest = image_hash(path)
        if digest in seen:
            repeated.append((image, seen[digest]))
            continue
        if len(selected) < 2:
            seen[digest] = item["row_id"]
            selected.append(image)
    if selected:
        nested = cell.add_table(rows=1, cols=len(selected))
        nested.autofit = False
        for image_cell, image in zip(nested.rows[0].cells, selected):
            clear_cell(image_cell)
            p = image_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(image["path"]), width=Inches(1.18))
            caption = add_cell_paragraph(
                image_cell,
                compact(image.get("label") or "证据截图", 35),
                size=6.1,
                color=MUTED,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                after=0,
            )
            add_hyperlink(caption, "原图", as_uri(str(image["path"])))
    if repeated:
        refs = sorted({row_id for _, row_id in repeated})
        p = add_cell_paragraph(
            cell,
            "重复图片未再次嵌入，首次出现于：" + "、".join(refs),
            size=6.4,
            color=MUTED,
            after=0,
        )
        first = repeated[0][0]
        add_hyperlink(p, "打开原图", as_uri(str(first["path"])))
    return len(selected)


def add_main_table(doc: Document, summary: dict[str, Any]) -> int:
    heading = add_text(
        doc,
        "61项证据明细｜按状态决定是否需要你行动",
        size=16,
        color=NAVY,
        bold=True,
        after=2,
    )
    heading.paragraph_format.page_break_before = True
    add_text(
        doc,
        "主状态为“待处理”的行请按“项目经理怎么做”推进；主状态为“已查看”的行可先阅读AI结论，案件关键时再补强。",
        size=8,
        color=MUTED,
        after=5,
    )
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    set_table_fixed(table, MAIN_WIDTHS_CM)
        headers = (
            "状态",
            "编号｜搜证类别",
            "AI看过 / 没看成的平台",
            "AI结论与小图",
            "项目经理怎么做",
            "要保存什么",
        )
    for cell, header in zip(table.rows[0].cells, headers):
        clear_cell(cell)
        cell_shading(cell, HEADER_FILL)
        add_cell_paragraph(
            cell,
            header,
            size=7.3,
            color=NAVY,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    set_repeat_table_header(table.rows[0])
    embedded = 0
    seen_images: dict[str, str] = {}
    for item in summary["rows"]:
        row = table.add_row()
        set_table_fixed(table, MAIN_WIDTHS_CM)
        for cell in row.cells:
            clear_cell(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        style = STATUS[item["manager_status"]]
        cell_shading(row.cells[0], style["fill"])
        add_cell_paragraph(
            row.cells[0],
            item["status_label"].replace("｜", "\n"),
            size=7.4,
            color=style["accent"],
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        total = int((item.get("action_counts") or {}).get("total", 0) or 0)
        checked = int((item.get("action_counts") or {}).get("AI_CHECKED_FOUND", 0) or 0) + int(
            (item.get("action_counts") or {}).get("AI_CHECKED_NO_RESULT", 0) or 0
        )
        add_cell_paragraph(
            row.cells[0],
            f"AI已查 {checked}/{total}",
            size=6.3,
            color=MUTED,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            after=0,
        )
        if item.get("is_mixed_checked_and_unresolved"):
            add_cell_paragraph(
                row.cells[0],
                "混合行",
                size=6.2,
                color=style["accent"],
                bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                after=0,
            )
        add_cell_paragraph(
            row.cells[1],
            item.get("row_id_label") or f"{item['row_id']}｜{item['objective']}",
            size=7.2,
            color=NAVY,
            bold=True,
            after=0,
        )
        add_scope_cell(row.cells[2], item)
        add_cell_paragraph(
            row.cells[3],
            item["finding"],
            size=7.0,
            color=INK,
            bold=False,
            after=2,
        )
        source_paragraph = add_cell_paragraph(row.cells[3], "", size=6.6, after=1)
        for index, source in enumerate(item.get("sources") or []):
            if index:
                run = source_paragraph.add_run("  ")
                set_font(run, size=6.6, color=MUTED)
            add_hyperlink(
                source_paragraph,
                compact(source.get("label") or "来源", 24),
                str(source.get("url") or ""),
            )
            if index >= 2:
                break
        embedded += add_images(row.cells[3], item, seen_images)
        add_cell_paragraph(
            row.cells[4],
            item["manager_action"],
            size=7.0,
            color=INK,
            after=1,
        )
        for platform in (item.get("pending_platforms") or [])[:2]:
            url = str(platform.get("url") or "")
            if url:
                p = add_cell_paragraph(row.cells[4], "", size=6.6, after=0)
                add_hyperlink(p, f"打开{platform.get('platform') or '平台'}", url)
        add_cell_paragraph(
            row.cells[5],
            item["save_items"],
            size=6.9,
            color=INK,
            after=0,
        )
    return embedded


def build_docx(
    summary: dict[str, Any],
    output: Path,
    title: str,
    rights_holder: str,
    defendant_product: str,
    source_url: str,
) -> dict[str, int]:
    doc = Document()
    configure_document(doc)
    add_title_block(doc, title, rights_holder, defendant_product, source_url)
    add_dashboard(doc, summary)
    add_pending_table(doc, summary)
    embedded = add_main_table(doc, summary)
    output.parent.mkdir(parents=True, exist_ok=True)
    fd, temp_name = tempfile.mkstemp(
        prefix=output.stem + ".", suffix=".tmp.docx", dir=str(output.parent)
    )
    os.close(fd)
    temp_path = Path(temp_name)
    try:
        doc.save(temp_path)
        audit = validate_docx(temp_path, summary, embedded)
        os.replace(temp_path, output)
        return audit
    finally:
        if temp_path.exists():
            temp_path.unlink()


def validate_docx(path: Path, summary: dict[str, Any], expected_images: int) -> dict[str, int]:
    doc = Document(path)
    if len(doc.tables) != 3:
        raise ValueError(f"顶层表格应为3个，实际{len(doc.tables)}个")
    pending_table = doc.tables[1]
    main_table = doc.tables[2]
    pending_count = summary["dashboard"]["manager_pending_rows"]
    if len(pending_table.rows) != pending_count + 1 or len(pending_table.columns) != 6:
        raise ValueError("待办表结构错误")
    if len(main_table.rows) != 62 or len(main_table.columns) != 6:
        raise ValueError("主表必须为61任务行+1表头、6列")
    main_labels = [row.cells[0].text.replace("\n", "｜") for row in main_table.rows[1:]]
    for status in STATUS_ORDER:
        expected = summary["dashboard"]["status_counts"][status]
        actual = sum(STATUS[status]["label"] in label for label in main_labels)
        if actual != expected:
            raise ValueError(f"{STATUS[status]['label']}计数错误：{actual}!={expected}")
    with zipfile.ZipFile(path) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        document_xml = archive.read("word/document.xml")
        rels_xml = archive.read("word/_rels/document.xml.rels")
    if len(media) != expected_images:
        raise ValueError(f"内嵌图片数量错误：media={len(media)}, expected={expected_images}")
    hyperlink_nodes = document_xml.count(b"<w:hyperlink")
    external_links = rels_xml.count(b'TargetMode="External"')
    if hyperlink_nodes < 61:
        raise ValueError(f"超链接不足：{hyperlink_nodes}")
    return {
        "top_level_tables": len(doc.tables),
        "pending_rows": len(pending_table.rows) - 1,
        "main_rows": len(main_table.rows) - 1,
        "main_columns": len(main_table.columns),
        "embedded_images": len(media),
        "hyperlink_nodes": hyperlink_nodes,
        "external_relationships": external_links,
    }


def main() -> int:
    args = parse_args()
    snapshot = load_json(args.input)
    summary = build_summary(snapshot)
    summary_path = Path(args.summary_output)
    summary_path.parent.mkdir(parents=True, exist_ok=True)
    summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2))
    audit = build_docx(
        summary,
        Path(args.output),
        args.title,
        args.rights_holder,
        args.defendant_product,
        args.source_url,
    )
    print(
        json.dumps(
            {
                "ok": True,
                "output": str(Path(args.output)),
                "summary": str(summary_path),
                "dashboard": summary["dashboard"],
                "audit": audit,
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
