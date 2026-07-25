#!/usr/bin/env python3
"""Render the 61-row heavy-case checklist as one compact inline-evidence table."""

from __future__ import annotations

import argparse
import hashlib
import io
import json
import os
import re
import sys
import uuid
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Pt, RGBColor
from lxml import etree


TASK_GROUPS: Sequence[Tuple[str, Sequence[Tuple[str, str]]]] = (
    (
        "一、侵权规模与获利能力",
        (
            ("1-1-1", "测购确认侵权主体"),
            ("1-1-2", "工商登记与经营能力"),
            ("1-1-3", "展会身份与产品线索"),
            ("1-1-4", "官网及自媒体主体线索"),
            ("1-1-5", "产品备案、认证与标准"),
            ("1-1-6", "招投标供应主体"),
            ("1-1-7", "专利申请人及生产主体"),
            ("1-1-8", "海关进出口主体"),
            ("1-2-1", "短视频及内容宣传渠道"),
            ("1-2-2", "本地生活及旅游渠道"),
            ("1-2-3", "招投标销售渠道"),
            ("1-2-4", "电商平台销售渠道"),
            ("1-2-5", "共享、租赁、加盟及连锁渠道"),
            ("1-3-1", "环评、投资及产线规模"),
            ("1-3-2", "招投标数量与金额"),
            ("1-3-3", "政府批文、项目产值"),
            ("1-3-4", "行业份额、产量与销量"),
            ("1-3-5", "招股书、年报及集团披露"),
            ("1-3-6", "注册资本及实缴情况"),
            ("1-4-1", "国内销售与发货区域"),
            ("1-4-2", "海外销售与进出口区域"),
            ("1-5-1", "中标数量与金额"),
            ("1-5-2", "年报销量与营业收入"),
            ("1-5-3", "第三方销量数据"),
            ("1-5-4", "被告自认销量与产能"),
            ("1-5-5", "线上多店销量汇总"),
        ),
    ),
    (
        "二、权利价值与贡献度",
        (
            ("2-1-1", "被告毛利率与净利率"),
            ("2-1-2", "同类企业及行业利润率"),
            ("2-2-1", "权利、技术奖项与认可"),
            ("2-2-2", "权利许可使用费"),
            ("2-2-3", "专家意见"),
            ("2-2-4", "市场调查与问卷"),
            ("2-2-5", "消费者混淆或关注点评论"),
            ("2-2-6", "多模型相似性观察"),
            ("2-2-7", "权利贡献相关论文"),
            ("2-2-8", "横向及纵向对照"),
            ("2-3-1", "广告投入及代言宣传"),
            ("2-3-2", "消费者评价"),
            ("2-3-3", "新闻报道"),
            ("2-3-4", "原告销量、门店及用户覆盖"),
            ("2-3-5", "行业排名与奖项"),
        ),
    ),
    (
        "三、侵权恶意与情节",
        (
            ("3-1-1", "商品或包装权利标识"),
            ("3-1-2", "权利类型与可识别性"),
            ("3-1-3", "知名度、接触可能与主观恶意"),
            ("3-1-4", "多项权利及复合侵权线索"),
            ("3-1-5", "重复侵权记录"),
            ("3-1-6", "商标囤积及抢注"),
            ("3-2-1", "既往代理、合作或接触关系"),
            ("3-2-2", "任职或雇佣关系"),
            ("3-2-3", "既往诉讼、和解或承诺"),
            ("3-2-4", "行政处罚记录"),
            ("3-2-5", "在先警告或通知"),
            ("3-2-6", "多枚模仿商标"),
        ),
    ),
    (
        "四、裁判尺度与合理费用",
        (
            ("4-1", "同一权利关联案件及判赔"),
            ("4-2", "合理维权费用"),
        ),
    ),
    (
        "五、其他酌定因素与履行能力",
        (
            ("5-1", "侵权地区经济状况"),
            ("5-2-1", "残疾等法定或酌定因素"),
            ("5-2-2", "年龄及经营者个人状况"),
            ("5-2-3", "精神状态等特殊因素"),
            ("5-3-1", "被告资产及履行能力"),
            ("5-3-2", "失信、执行及破产信息"),
        ),
    ),
)

TASKS: Sequence[Tuple[str, str]] = tuple(
    task for _group_name, group_tasks in TASK_GROUPS for task in group_tasks
)
TASK_IDS: Sequence[str] = tuple(task_id for task_id, _title in TASKS)

PAGE_WIDTH_CM = 29.7
PAGE_HEIGHT_CM = 21.0
MARGIN_LEFT_RIGHT_CM = 0.65
MARGIN_TOP_BOTTOM_CM = 0.60
HEADER_FOOTER_CM = 0.30
COLUMN_WIDTHS_CM: Sequence[float] = (1.35, 5.15, 4.55, 12.25, 5.10)
TABLE_WIDTH_CM = sum(COLUMN_WIDTHS_CM)
TABLE_INDENT_DXA = 0
CELL_MARGINS_DXA = {"top": 45, "bottom": 45, "start": 55, "end": 55}

FONT_LATIN = "Aptos"
FONT_CJK = "PingFang SC"
COLOR_INK = "17202A"
COLOR_MUTED = "5D6D7E"
COLOR_NAVY = "17365D"
COLOR_TEAL = "0F6B67"
COLOR_HEADER_FILL = "DCE8F1"
COLOR_GROUP_FILL = "EAF3F1"
COLOR_BORDER = "AAB7C4"
COLOR_LINK = "0563C1"
COLOR_WHITE = "FFFFFF"

STATUS_META: Dict[str, Tuple[str, str, str]] = {
    "COMPLETE_VERIFIED": ("已核实", "0F6B67", "E7F3F1"),
    "PARTIAL_VERIFIED": ("部分核实", "8A5A00", "FFF5D8"),
    "LEAD_ONLY": ("有线索", "1D5F91", "E7F0F8"),
    "NOT_FOUND": ("未发现", "68727D", "F0F2F4"),
    "BLOCKED": ("受阻", "8B2E2E", "FBEAEA"),
    "NEEDS_HUMAN": ("需人工", "7D4F00", "FFF0D1"),
    "NOT_APPLICABLE": ("不适用", "68727D", "F0F2F4"),
    "ERROR": ("异常", "8B2E2E", "FBEAEA"),
}

URL_RE = re.compile(r"https?://[^\s<>\]\[()（）]+", re.IGNORECASE)
TIMESTAMP_RE = re.compile(
    r"\b20\d{2}[-/.]\d{1,2}[-/.]\d{1,2}(?:[T\s]+\d{1,2}:\d{2}(?::\d{2})?(?:\.\d+)?(?:Z|[+-]\d{2}:?\d{2})?)?\b"
)
LONG_HEX_RE = re.compile(r"\b[a-fA-F0-9]{32,}\b")
WHITESPACE_RE = re.compile(r"\s+")


@dataclass(frozen=True)
class PreparedImage:
    path: Path
    role: str
    label: str
    source_url: str
    original_url: str
    sha256: str
    pixel_width: int
    pixel_height: int


@dataclass
class RowMedia:
    display: List[PreparedImage] = field(default_factory=list)
    duplicate_rows: List[str] = field(default_factory=list)
    extras: List[PreparedImage] = field(default_factory=list)
    missing_count: int = 0
    unsupported_count: int = 0
    declared_count: int = 0


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Render a compact A4-landscape, five-column heavy-case evidence report."
    )
    parser.add_argument("--results", required=True, type=Path)
    parser.add_argument("--attachments", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--run-dir", required=True, type=Path)
    parser.add_argument(
        "--allow-partial",
        action="store_true",
        help="Allow a report while READY/RETRYABLE/RUNNING actions remain and mark it PARTIAL.",
    )
    parser.add_argument("--title", default="重案调查一般证据清单")
    parser.add_argument("--rights-holder", default="")
    parser.add_argument("--defendant-product", default="")
    parser.add_argument("--store-name", default="")
    parser.add_argument("--operator-name", default="")
    parser.add_argument("--source-url", default="")
    parser.add_argument("--source-label", default="种子商品链接")
    parser.add_argument("--investigation-date", default="")
    return parser.parse_args()


def cm_to_dxa(value: float) -> int:
    return int(round(value / 2.54 * 1440))


def load_json(path: Path) -> Any:
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def load_run_gate(run_dir: Path, allow_partial: bool) -> Dict[str, Any]:
    resolved = run_dir.expanduser().resolve()
    state = load_json(resolved / "run_state.json")
    queue = load_json(resolved / "action_queue.json")
    if state.get("run_id") != queue.get("run_id") or state.get("case_id") != queue.get("case_id"):
        raise ValueError("run_state and action_queue identifiers do not match")
    actions = queue.get("actions")
    if not isinstance(actions, list):
        raise ValueError("action_queue.actions must be an array")
    active_states = ("READY", "RETRYABLE", "RUNNING")
    actual = {
        name: sum(
            1 for action in actions
            if isinstance(action, dict) and action.get("execution_state") == name
        )
        for name in active_states
    }
    declared = state.get("counts", {}).get("by_execution_state", {})
    for name in active_states:
        try:
            actual[name] = max(actual[name], int(declared.get(name, 0)))
        except (TypeError, ValueError):
            raise ValueError(f"run_state count for {name} is invalid")
    top_state = str(state.get("execution_state") or "")
    if top_state in active_states:
        actual[top_state] = max(actual[top_state], 1)
    active_count = sum(actual.values())
    if active_count and not allow_partial:
        raise ValueError(
            "formal report blocked: READY/RETRYABLE/RUNNING actions remain; "
            "use --allow-partial only for an explicitly marked partial report"
        )
    return {
        "run_id": state.get("run_id"),
        "case_state": state.get("case_state"),
        "execution_state": state.get("execution_state"),
        "active_counts": actual,
        "active_count": active_count,
        "partial": bool(allow_partial),
    }


def unwrap_rows(payload: Any, keys: Sequence[str], label: str) -> List[Dict[str, Any]]:
    if isinstance(payload, list):
        rows = payload
    elif isinstance(payload, dict):
        rows = None
        for key in keys:
            candidate = payload.get(key)
            if isinstance(candidate, list):
                rows = candidate
                break
        if rows is None:
            raise ValueError(f"{label} JSON must be an array or contain one of: {', '.join(keys)}")
    else:
        raise ValueError(f"{label} JSON has unsupported top-level type")
    if not all(isinstance(row, dict) for row in rows):
        raise ValueError(f"{label} rows must be JSON objects")
    return rows


def row_identifier(row: Dict[str, Any]) -> str:
    for key in ("item_id", "row_id", "task_id", "id"):
        value = row.get(key)
        if value is not None and str(value).strip():
            return str(value).strip()
    return ""


def index_rows(rows: Iterable[Dict[str, Any]], label: str) -> Dict[str, Dict[str, Any]]:
    indexed: Dict[str, Dict[str, Any]] = {}
    for row in rows:
        row_id = row_identifier(row)
        if not row_id:
            raise ValueError(f"{label} row is missing item_id/row_id/task_id")
        if row_id in indexed:
            raise ValueError(f"{label} contains duplicate row id: {row_id}")
        indexed[row_id] = row
    return indexed


def load_inputs(results_path: Path, attachments_path: Path) -> Tuple[Dict[str, Dict[str, Any]], Dict[str, Dict[str, Any]]]:
    result_rows = unwrap_rows(load_json(results_path), ("results", "items", "rows"), "results")
    attachment_rows = unwrap_rows(
        load_json(attachments_path), ("attachments", "items", "rows"), "attachments"
    )
    results = index_rows(result_rows, "results")
    attachments = index_rows(attachment_rows, "attachments")

    expected = set(TASK_IDS)
    result_ids = set(results)
    if result_ids != expected:
        missing = sorted(expected - result_ids)
        extra = sorted(result_ids - expected)
        raise ValueError(
            "results must contain exactly the 61 canonical task rows; "
            f"missing={missing or 'none'}, extra={extra or 'none'}"
        )
    unknown_attachments = sorted(set(attachments) - expected)
    if unknown_attachments:
        raise ValueError(f"attachments contains unknown row ids: {unknown_attachments}")
    return results, attachments


def sanitize_text(value: Any, limit: int) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        value = json.dumps(value, ensure_ascii=False, separators=(",", ":"))
    text = str(value)
    text = URL_RE.sub("来源链接", text)
    text = TIMESTAMP_RE.sub("", text)
    text = LONG_HEX_RE.sub("", text)
    text = text.replace("\r", " ").replace("\n", " ").replace("\t", " ")
    text = WHITESPACE_RE.sub(" ", text).strip(" ;；,，。|｜")
    text = re.sub(r"([；;，,。])\1+", r"\1", text)
    if len(text) > limit:
        text = text[: max(1, limit - 1)].rstrip(" ;；,，。") + "…"
    return text


def text_list(value: Any) -> List[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item) for item in value if str(item).strip()]
    if isinstance(value, str) and value.strip():
        return [value]
    return []


def is_checked(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in {"1", "true", "yes", "y", "checked", "done", "complete"}


def is_web_url(value: Any) -> bool:
    return isinstance(value, str) and value.lower().startswith(("https://", "http://"))


def resolve_image_path(raw_path: Any, attachments_path: Path) -> Optional[Path]:
    if not raw_path or not str(raw_path).strip():
        return None
    candidate = Path(str(raw_path)).expanduser()
    candidates = [candidate] if candidate.is_absolute() else [
        attachments_path.parent / candidate,
        attachments_path.parent.parent / candidate,
        Path.cwd() / candidate,
    ]
    for path in candidates:
        if path.is_file():
            return path.resolve()
    return candidates[0].resolve()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def image_dimensions(path: Path) -> Optional[Tuple[int, int]]:
    try:
        with Image.open(path) as image:
            image.verify()
        with Image.open(path) as image:
            return int(image.width), int(image.height)
    except Exception:
        return None


def source_url_for_image(
    image: Dict[str, Any], attachment: Dict[str, Any], result: Dict[str, Any]
) -> str:
    direct = image.get("source_url") or image.get("url")
    if is_web_url(direct):
        return str(direct)

    role = str(image.get("role") or "").strip().lower()
    attachment_sources = attachment.get("source_urls") or []
    result_sources = result.get("sources") or []
    sources = [source for source in attachment_sources + result_sources if isinstance(source, dict)]
    for source in sources:
        source_role = str(source.get("role") or source.get("publisher") or "").strip().lower()
        url = source.get("url")
        if role and role == source_role and is_web_url(url):
            return str(url)
    for source in sources:
        url = source.get("url")
        if is_web_url(url):
            return str(url)
    return ""


def prepare_media(
    results: Dict[str, Dict[str, Any]],
    attachments: Dict[str, Dict[str, Any]],
    attachments_path: Path,
) -> Tuple[Dict[str, RowMedia], Dict[str, Any]]:
    primary_by_sha: Dict[str, str] = {}
    prepared: Dict[str, RowMedia] = {}
    total_declared = 0
    total_displayed = 0
    duplicate_occurrences = 0

    for row_id in TASK_IDS:
        attachment = attachments.get(row_id, {})
        result = results[row_id]
        raw_images = attachment.get("images") or []
        if not isinstance(raw_images, list):
            raw_images = []
        media = RowMedia(declared_count=len(raw_images))
        total_declared += len(raw_images)
        seen_this_row = set()

        for raw_image in raw_images:
            if not isinstance(raw_image, dict):
                media.unsupported_count += 1
                continue
            path = resolve_image_path(raw_image.get("path"), attachments_path)
            if path is None or not path.is_file():
                media.missing_count += 1
                continue
            dimensions = image_dimensions(path)
            if dimensions is None:
                media.unsupported_count += 1
                continue
            digest = sha256_file(path)
            if digest in seen_this_row:
                duplicate_occurrences += 1
                continue
            seen_this_row.add(digest)
            if digest in primary_by_sha:
                primary_row = primary_by_sha[digest]
                if primary_row != row_id and primary_row not in media.duplicate_rows:
                    media.duplicate_rows.append(primary_row)
                duplicate_occurrences += 1
                continue

            item = PreparedImage(
                path=path,
                role=sanitize_text(raw_image.get("role") or "证据", 16),
                label=sanitize_text(raw_image.get("label") or path.stem, 34),
                source_url=source_url_for_image(raw_image, attachment, result),
                original_url=path.as_uri(),
                sha256=digest,
                pixel_width=dimensions[0],
                pixel_height=dimensions[1],
            )
            if len(media.display) < 2:
                media.display.append(item)
                primary_by_sha[digest] = row_id
                total_displayed += 1
            else:
                media.extras.append(item)

        prepared[row_id] = media

    stats = {
        "declared_images": total_declared,
        "displayed_unique_images": total_displayed,
        "duplicate_occurrences": duplicate_occurrences,
        "unique_sha_groups": len(primary_by_sha),
        "missing_images": sum(item.missing_count for item in prepared.values()),
        "unsupported_images": sum(item.unsupported_count for item in prepared.values()),
        "extra_images_not_thumbnailed": sum(len(item.extras) for item in prepared.values()),
    }
    return prepared, stats


def set_run_font(
    run: Any,
    size: float,
    color: str = COLOR_INK,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = FONT_LATIN
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_CJK)


def format_paragraph(
    paragraph: Any,
    size: float = 7.0,
    align: Optional[Any] = None,
    before: float = 0,
    after: float = 0,
    line_spacing: float = 1.0,
    keep_with_next: bool = False,
) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.widow_control = False
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        set_run_font(run, size)


def clear_paragraph(paragraph: Any) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_external_hyperlink(
    paragraph: Any,
    text: str,
    target: str,
    size: float = 6.3,
    bold: bool = False,
    color: str = COLOR_LINK,
) -> None:
    if not target:
        run = paragraph.add_run(text)
        set_run_font(run, size, color=color, bold=bold)
        return
    relationship_id = paragraph.part.relate_to(target, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_CJK)
    rpr.append(rfonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    rpr.append(color_el)
    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), str(int(round(size * 2))))
    rpr.append(size_el)
    size_cs_el = OxmlElement("w:szCs")
    size_cs_el.set(qn("w:val"), str(int(round(size * 2))))
    rpr.append(size_cs_el)
    if bold:
        rpr.append(OxmlElement("w:b"))
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_internal_hyperlink(
    paragraph: Any, text: str, anchor: str, size: float = 6.3
) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_CJK)
    rpr.append(rfonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), COLOR_LINK)
    rpr.append(color_el)
    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), str(int(round(size * 2))))
    rpr.append(size_el)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def bookmark_name(row_id: str) -> str:
    return "row_" + re.sub(r"[^A-Za-z0-9_]", "_", row_id)


def add_bookmark(paragraph: Any, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def picture_size(image: PreparedImage, image_count: int) -> Tuple[float, float]:
    # Named compactness override: intentionally below the contractual maxima.
    if image_count == 1:
        max_width, max_height = 5.20, 2.90
    else:
        max_width, max_height = 4.55, 2.55
    aspect = image.pixel_width / max(1, image.pixel_height)
    width = min(max_width, max_height * aspect)
    height = width / max(aspect, 0.01)
    if height > max_height:
        height = max_height
        width = height * aspect
    return width, height


def picture_source(path: Path) -> Any:
    if path.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff"}:
        return str(path)
    stream = io.BytesIO()
    with Image.open(path) as image:
        if image.mode not in {"RGB", "RGBA"}:
            image = image.convert("RGB")
        image.save(stream, format="PNG")
    stream.seek(0)
    return stream


def add_hyperlinked_picture(
    paragraph: Any,
    image: PreparedImage,
    width_cm: float,
    height_cm: float,
) -> None:
    run = paragraph.add_run()
    inline_shape = run.add_picture(
        picture_source(image.path), width=Cm(width_cm), height=Cm(height_cm)
    )
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", sanitize_text(image.label, 80))
    target = image.source_url or image.original_url
    relationship_id = paragraph.part.relate_to(target, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run_element = run._r
    paragraph._p.remove(run_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def set_cell_width(cell: Any, width_cm: float) -> None:
    width_dxa = cm_to_dxa(width_cm)
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGINS_DXA.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell: Any, color: str = COLOR_BORDER, size: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table: Any) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_pr = table._tbl.tblPr
    table_width = table_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(cm_to_dxa(TABLE_WIDTH_CM)))
    table_width.set(qn("w:type"), "dxa")

    table_indent = table_pr.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_pr.append(table_indent)
    table_indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    table_indent.set(qn("w:type"), "dxa")

    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
    for grid_column, width_cm in zip(grid_columns, COLUMN_WIDTHS_CM):
        grid_column.set(qn("w:w"), str(cm_to_dxa(width_cm)))


def set_row_repeat(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def set_row_cant_split(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def add_field(paragraph: Any, field_name: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" {field_name} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))
    set_run_font(run, 6.5, color=COLOR_MUTED)


def configure_document(document: Any) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.header_distance = Cm(HEADER_FOOTER_CM)
    section.footer_distance = Cm(HEADER_FOOTER_CM)

    normal = document.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal.font.size = Pt(7)
    normal.font.color.rgb = RGBColor.from_string(COLOR_INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    header = section.header.paragraphs[0]
    clear_paragraph(header)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run("重案证据调查 · 一般证据清单")
    set_run_font(header_run, 6.5, color=COLOR_MUTED, bold=True)
    format_paragraph(header, after=0)

    footer = section.footer.paragraphs[0]
    clear_paragraph(footer)
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = footer.add_run("第 ")
    set_run_font(prefix, 6.5, color=COLOR_MUTED)
    add_field(footer, "PAGE")
    slash = footer.add_run(" / ")
    set_run_font(slash, 6.5, color=COLOR_MUTED)
    add_field(footer, "NUMPAGES")
    suffix = footer.add_run(" 页")
    set_run_font(suffix, 6.5, color=COLOR_MUTED)
    format_paragraph(footer, after=0)


def add_title_block(document: Any, args: argparse.Namespace) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run(sanitize_text(args.title, 80))
    set_run_font(title_run, 13, color=COLOR_NAVY, bold=True)
    format_paragraph(title, after=1.5, line_spacing=1.0, keep_with_next=True)

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.LEFT
    parts = []
    if args.rights_holder:
        parts.append(("权利人", sanitize_text(args.rights_holder, 36)))
    if args.defendant_product:
        parts.append(("被告商品", sanitize_text(args.defendant_product, 64)))
    if args.store_name:
        parts.append(("店铺", sanitize_text(args.store_name, 36)))
    if args.operator_name:
        parts.append(("经营主体", sanitize_text(args.operator_name, 46)))
    if args.investigation_date:
        parts.append(("调查日", sanitize_text(args.investigation_date, 20)))
    for index, (label, value) in enumerate(parts):
        if index:
            separator = metadata.add_run("   |   ")
            set_run_font(separator, 6.8, color=COLOR_MUTED)
        label_run = metadata.add_run(f"{label}：")
        set_run_font(label_run, 6.8, color=COLOR_MUTED, bold=True)
        value_run = metadata.add_run(value)
        set_run_font(value_run, 6.8, color=COLOR_INK)
    if args.source_url:
        if parts:
            separator = metadata.add_run("   |   ")
            set_run_font(separator, 6.8, color=COLOR_MUTED)
        add_external_hyperlink(
            metadata,
            sanitize_text(args.source_label, 24),
            args.source_url,
            size=6.8,
            bold=True,
        )
    if not parts and not args.source_url:
        run = metadata.add_run("案件信息以 results 与 attachments 数据为准")
        set_run_font(run, 6.8, color=COLOR_MUTED)
    format_paragraph(metadata, after=2.5, line_spacing=1.0, keep_with_next=True)


def configure_cell(cell: Any, width_cm: float, vertical_top: bool = True) -> None:
    set_cell_width(cell, width_cm)
    set_cell_margins(cell)
    set_cell_border(cell)
    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.TOP if vertical_top else WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )


def add_cell_run(
    paragraph: Any,
    text: str,
    size: float,
    color: str = COLOR_INK,
    bold: bool = False,
    italic: bool = False,
) -> Any:
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return run


def fill_header_row(table: Any) -> None:
    labels = ("编号", "证据项目", "已查看平台 / 网站", "调查结果与行内证据", "下一步重点")
    row = table.rows[0]
    set_row_repeat(row)
    set_row_cant_split(row)
    for index, (cell, label, width_cm) in enumerate(zip(row.cells, labels, COLUMN_WIDTHS_CM)):
        configure_cell(cell, width_cm, vertical_top=False)
        shade_cell(cell, COLOR_HEADER_FILL)
        paragraph = cell.paragraphs[0]
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_cell_run(paragraph, label, 7.2, color=COLOR_NAVY, bold=True)
        format_paragraph(paragraph, after=0, line_spacing=1.0)


def add_group_row(table: Any, label: str) -> None:
    row = table.add_row()
    set_row_cant_split(row)
    merged = row.cells[0].merge(row.cells[-1])
    configure_cell(merged, TABLE_WIDTH_CM, vertical_top=False)
    shade_cell(merged, COLOR_GROUP_FILL)
    paragraph = merged.paragraphs[0]
    clear_paragraph(paragraph)
    add_cell_run(paragraph, label, 7.4, color=COLOR_TEAL, bold=True)
    format_paragraph(paragraph, after=0, line_spacing=1.0, keep_with_next=True)


def platform_entries(result: Dict[str, Any]) -> List[Tuple[str, bool, str]]:
    raw_entries = result.get("checked_platforms") or []
    if not isinstance(raw_entries, list):
        return []
    entries = []
    seen = set()
    for raw in raw_entries:
        if not isinstance(raw, dict):
            continue
        name = sanitize_text(raw.get("name") or raw.get("platform") or raw.get("site"), 24)
        if not name:
            continue
        key = name.casefold()
        if key in seen:
            continue
        seen.add(key)
        url = raw.get("url") if is_web_url(raw.get("url")) else ""
        entries.append((name, is_checked(raw.get("checked")), str(url)))
    return entries


def fill_platform_cell(cell: Any, result: Dict[str, Any]) -> None:
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    entries = platform_entries(result)
    if not entries:
        add_cell_run(paragraph, "☒ 未记录平台", 6.5, color=COLOR_MUTED)
    for index, (name, checked, url) in enumerate(entries):
        if index:
            add_cell_run(paragraph, "  ", 6.2, color=COLOR_MUTED)
        marker = "☑ " if checked else "☒ "
        add_cell_run(
            paragraph,
            marker,
            6.6,
            color=COLOR_TEAL if checked else COLOR_MUTED,
            bold=True,
        )
        if url:
            add_external_hyperlink(paragraph, name, url, size=6.5)
        else:
            add_cell_run(paragraph, name, 6.5)
    format_paragraph(paragraph, after=0, line_spacing=1.0)


def status_meta(result: Dict[str, Any]) -> Tuple[str, str, str, str]:
    status = str(result.get("status") or "LEAD_ONLY").strip().upper()
    label, color, fill = STATUS_META.get(status, (status or "待核验", "68727D", "F0F2F4"))
    return status, label, color, fill


def conclusion_text(result: Dict[str, Any]) -> str:
    for key in ("compact_note", "reason", "findings"):
        value = result.get(key)
        if value:
            text = sanitize_text(value, 118)
            if text:
                return text
    return "本轮未形成可直接写入的公开调查结论。"


def no_image_reason(status: str, media: RowMedia) -> str:
    if media.duplicate_rows:
        return "去重｜同图已在前行展示"
    if media.missing_count:
        return "无图｜附件文件缺失"
    if media.unsupported_count:
        return "无图｜附件格式不可渲染"
    mapping = {
        "NOT_FOUND": "无图｜公开查询无结果",
        "BLOCKED": "无图｜登录或技术受阻",
        "ERROR": "无图｜查询异常",
        "NEEDS_HUMAN": "无图｜需线下或人工完成",
        "NOT_APPLICABLE": "无图｜本项不适用",
        "PARTIAL_VERIFIED": "无图｜已有结论，待补截图",
        "LEAD_ONLY": "无图｜仅有线索，待补截图",
        "COMPLETE_VERIFIED": "无图｜结论已核实，未附截图",
    }
    return mapping.get(status, "无图｜待补可视证据")


def add_media_content(cell: Any, result: Dict[str, Any], media: RowMedia) -> None:
    status, status_label, status_color, status_fill = status_meta(result)
    shade_cell(cell, status_fill)
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    add_cell_run(paragraph, f"【{status_label}】", 6.8, color=status_color, bold=True)
    add_cell_run(paragraph, conclusion_text(result), 6.8, color=COLOR_INK)
    format_paragraph(paragraph, after=0.6 if media.display else 0, line_spacing=1.0)

    if media.display:
        image_paragraph = cell.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for index, item in enumerate(media.display):
            if index:
                spacer = image_paragraph.add_run("  ")
                set_run_font(spacer, 5.5, color=COLOR_MUTED)
            width_cm, height_cm = picture_size(item, len(media.display))
            add_hyperlinked_picture(image_paragraph, item, width_cm, height_cm)
        format_paragraph(image_paragraph, after=0, line_spacing=1.0)

        links = cell.add_paragraph()
        links.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for index, item in enumerate(media.display):
            if index:
                add_cell_run(links, "   |   ", 6.0, color=COLOR_MUTED)
            caption = sanitize_text(item.label or item.role or f"图{index + 1}", 22)
            add_cell_run(links, f"图{index + 1} {caption} ", 6.0, color=COLOR_MUTED)
            add_external_hyperlink(links, "原图", item.original_url, size=6.0)
            if item.source_url:
                add_cell_run(links, "/", 6.0, color=COLOR_MUTED)
                add_external_hyperlink(links, "来源", item.source_url, size=6.0)
        format_paragraph(links, after=0, line_spacing=1.0)
    else:
        reason_paragraph = cell.add_paragraph()
        add_cell_run(
            reason_paragraph,
            no_image_reason(status, media),
            6.2,
            color=COLOR_MUTED,
            bold=True,
        )
        format_paragraph(reason_paragraph, after=0, line_spacing=1.0)

    if media.duplicate_rows:
        duplicate_paragraph = cell.add_paragraph()
        add_cell_run(duplicate_paragraph, "重复图：", 6.2, color=COLOR_MUTED, bold=True)
        for index, primary_row in enumerate(media.duplicate_rows):
            if index:
                add_cell_run(duplicate_paragraph, "、", 6.2, color=COLOR_MUTED)
            add_internal_hyperlink(
                duplicate_paragraph,
                f"同图见 {primary_row}",
                bookmark_name(primary_row),
                size=6.2,
            )
        format_paragraph(duplicate_paragraph, after=0, line_spacing=1.0)

    if media.extras:
        extra_paragraph = cell.add_paragraph()
        add_cell_run(
            extra_paragraph,
            f"另有 {len(media.extras)} 张原始附件：",
            6.0,
            color=COLOR_MUTED,
        )
        for index, item in enumerate(media.extras[:3]):
            if index:
                add_cell_run(extra_paragraph, "、", 6.0, color=COLOR_MUTED)
            add_external_hyperlink(extra_paragraph, f"原图{index + 1}", item.original_url, size=6.0)
        if len(media.extras) > 3:
            add_cell_run(extra_paragraph, "等", 6.0, color=COLOR_MUTED)
        format_paragraph(extra_paragraph, after=0, line_spacing=1.0)


def fill_next_steps(cell: Any, result: Dict[str, Any]) -> None:
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    focus_rows = result.get("next_step_focus")
    steps: List[str] = []
    if isinstance(focus_rows, list):
        priority_order = {"P0": 0, "P1": 1, "P2": 2, "P3": 3}
        ordered = sorted(
            (item for item in focus_rows if isinstance(item, dict)),
            key=lambda item: priority_order.get(str(item.get("priority") or "P3"), 9),
        )
        for item in ordered[:2]:
            priority = sanitize_text(item.get("priority") or "P3", 4)
            owner = sanitize_text(item.get("owner"), 10)
            action = sanitize_text(item.get("action"), 64)
            if action:
                steps.append("｜".join(part for part in (priority, owner, action) if part))
    if not steps:
        steps = [sanitize_text(step, 72) for step in text_list(result.get("next_steps"))]
        steps = [step for step in steps if step][:2]
    if not steps:
        status = str(result.get("status") or "").strip().upper()
        steps = [
            "本阶段无需补查。"
            if status in {"COMPLETE_VERIFIED", "NOT_APPLICABLE"}
            else "围绕本行证据缺口继续核验。"
        ]
    for index, step in enumerate(steps):
        target = paragraph if index == 0 else cell.add_paragraph()
        label = "优先｜" if index == 0 else "补强｜"
        add_cell_run(target, label, 6.5, color=COLOR_TEAL, bold=True)
        add_cell_run(target, step, 6.5, color=COLOR_INK)
        format_paragraph(target, after=0, line_spacing=1.0)


def add_task_row(
    table: Any,
    row_id: str,
    title: str,
    result: Dict[str, Any],
    media: RowMedia,
    bookmark_id: int,
) -> None:
    row = table.add_row()
    set_row_cant_split(row)
    cells = row.cells
    for cell, width_cm in zip(cells, COLUMN_WIDTHS_CM):
        configure_cell(cell, width_cm)

    id_paragraph = cells[0].paragraphs[0]
    clear_paragraph(id_paragraph)
    id_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_cell_run(id_paragraph, row_id, 6.8, color=COLOR_NAVY, bold=True)
    add_bookmark(id_paragraph, bookmark_name(row_id), bookmark_id)
    format_paragraph(id_paragraph, after=0, line_spacing=1.0)

    title_paragraph = cells[1].paragraphs[0]
    clear_paragraph(title_paragraph)
    add_cell_run(title_paragraph, title, 6.9, color=COLOR_INK, bold=True)
    format_paragraph(title_paragraph, after=0, line_spacing=1.0)

    fill_platform_cell(cells[2], result)
    add_media_content(cells[3], result, media)
    fill_next_steps(cells[4], result)


def build_document(
    args: argparse.Namespace,
    results: Dict[str, Dict[str, Any]],
    media_by_row: Dict[str, RowMedia],
) -> Any:
    document = Document()
    configure_document(document)
    document.core_properties.title = sanitize_text(args.title, 120)
    document.core_properties.subject = "61项重案证据调查清单"
    document.core_properties.author = "evidence-collection"
    document.core_properties.keywords = "知识产权, 证据清单, 被告产品, 重案调查"
    add_title_block(document, args)
    if getattr(args, "partial_report", False):
        warning = document.add_paragraph()
        warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = warning.add_run("PARTIAL / 部分输出：运行队列尚未耗尽，仅供阶段性审阅")
        set_run_font(run, 10.5, color="B42318", bold=True)
        format_paragraph(warning, after=4, line_spacing=1.0)

    table = document.add_table(rows=1, cols=5)
    set_table_geometry(table)
    fill_header_row(table)
    bookmark_id = 1000
    for group_name, tasks in TASK_GROUPS:
        add_group_row(table, group_name)
        for row_id, title in tasks:
            bookmark_id += 1
            add_task_row(
                table,
                row_id,
                title,
                results[row_id],
                media_by_row[row_id],
                bookmark_id,
            )
    return document


def validate_docx(path: Path) -> Dict[str, Any]:
    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    }
    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml")
    root = etree.fromstring(document_xml)

    tables = root.xpath(".//w:tbl", namespaces=namespaces)
    if len(tables) != 1:
        raise ValueError(f"structural audit failed: expected one table, got {len(tables)}")
    table = tables[0]
    grid_widths = [
        int(node.get(qn("w:w")))
        for node in table.xpath("./w:tblGrid/w:gridCol", namespaces=namespaces)
    ]
    expected_widths = [cm_to_dxa(value) for value in COLUMN_WIDTHS_CM]
    if grid_widths != expected_widths:
        raise ValueError(
            f"structural audit failed: grid widths {grid_widths} != {expected_widths}"
        )

    page_size = root.xpath(".//w:sectPr/w:pgSz[last()]", namespaces=namespaces)
    if not page_size:
        raise ValueError("structural audit failed: page size missing")
    page_width = int(page_size[0].get(qn("w:w")))
    page_height = int(page_size[0].get(qn("w:h")))
    orientation = page_size[0].get(qn("w:orient"))
    if page_width <= page_height or orientation != "landscape":
        raise ValueError("structural audit failed: page is not A4 landscape")

    bookmarks = root.xpath(".//w:bookmarkStart", namespaces=namespaces)
    task_bookmarks = {
        node.get(qn("w:name")) for node in bookmarks if (node.get(qn("w:name")) or "").startswith("row_")
    }
    expected_bookmarks = {bookmark_name(row_id) for row_id in TASK_IDS}
    if task_bookmarks != expected_bookmarks:
        raise ValueError("structural audit failed: canonical 61 row bookmarks are incomplete")

    drawings = root.xpath(".//w:drawing", namespaces=namespaces)
    anchors = root.xpath(".//wp:anchor", namespaces=namespaces)
    if anchors:
        raise ValueError("structural audit failed: floating image anchors are not allowed")

    data_rows = 0
    for table_row in table.xpath("./w:tr", namespaces=namespaces):
        row_bookmarks = table_row.xpath(".//w:bookmarkStart", namespaces=namespaces)
        names = [node.get(qn("w:name")) or "" for node in row_bookmarks]
        if not any(name.startswith("row_") for name in names):
            continue
        data_rows += 1
        cells = table_row.xpath("./w:tc", namespaces=namespaces)
        if len(cells) != 5:
            raise ValueError("structural audit failed: a task row does not have five cells")
        for index, cell in enumerate(cells):
            image_count = len(cell.xpath(".//w:drawing", namespaces=namespaces))
            if index != 3 and image_count:
                raise ValueError("structural audit failed: image found outside result column")
            if index == 3 and image_count > 2:
                raise ValueError("structural audit failed: more than two thumbnails in a row")
    if data_rows != 61:
        raise ValueError(f"structural audit failed: expected 61 task rows, got {data_rows}")

    visible_text = "".join(root.xpath(".//w:t/text()", namespaces=namespaces))
    if "证据截图附录" in visible_text:
        raise ValueError("structural audit failed: screenshot appendix is forbidden")
    if re.search(r"https?://", visible_text, re.IGNORECASE):
        raise ValueError("structural audit failed: a raw URL is visible in report text")

    hyperlinks = root.xpath(".//w:hyperlink", namespaces=namespaces)
    return {
        "task_rows": data_rows,
        "table_count": len(tables),
        "inline_images": len(drawings),
        "hyperlinks": len(hyperlinks),
        "grid_widths_dxa": grid_widths,
        "page_landscape": True,
    }


def atomic_save(document: Any, output: Path) -> Dict[str, Any]:
    output = output.expanduser().resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    temporary = output.with_name(
        f".{output.stem}.{os.getpid()}.{uuid.uuid4().hex}.tmp.docx"
    )
    try:
        document.save(temporary)
        audit = validate_docx(temporary)
        os.replace(temporary, output)
        return audit
    except Exception:
        try:
            temporary.unlink()
        except FileNotFoundError:
            pass
        raise


def main() -> int:
    args = parse_args()
    try:
        run_gate = load_run_gate(args.run_dir, args.allow_partial)
        args.partial_report = bool(args.allow_partial)
        results, attachments = load_inputs(args.results, args.attachments)
        media_by_row, media_stats = prepare_media(results, attachments, args.attachments)
        document = build_document(args, results, media_by_row)
        audit = atomic_save(document, args.output)
    except Exception as exc:
        print(f"render_compact_inline: {exc}", file=sys.stderr)
        return 1

    summary = {
        "output": str(args.output.expanduser().resolve()),
        **audit,
        **media_stats,
        "run_gate": run_gate,
    }
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
